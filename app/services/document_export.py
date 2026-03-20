from __future__ import annotations

import re
from pathlib import Path
from uuid import uuid4

from docx import Document

from app.core.config import settings


def _safe_name(value: str | None) -> str:
    if not value:
        return "document"
    cleaned = re.sub(r"[^\w\-.()\u4e00-\u9fff]+", "_", value).strip("._")
    return cleaned or "document"


def export_document_content(
    document_id: str,
    document_name: str | None,
    content: str,
) -> tuple[str, str]:
    export_dir = Path(settings.document_export_dir)
    if not export_dir.is_absolute():
        export_dir = Path.cwd() / export_dir
    export_dir.mkdir(parents=True, exist_ok=True)

    base_name = _safe_name(document_name)
    file_name = base_name if base_name.lower().endswith(".docx") else f"{base_name}.docx"
    file_path = export_dir / file_name
    if file_path.exists():
        stem = file_path.stem
        suffix = file_path.suffix
        file_name = f"{stem}_{uuid4().hex[:8]}{suffix}"
        file_path = export_dir / file_name

    document = Document()
    for paragraph in content.splitlines():
        document.add_paragraph(paragraph)
    if not content.splitlines():
        document.add_paragraph("")
    document.save(file_path)

    download_url = f"{settings.api_v1_prefix}/database/exports/{file_name}"
    return str(file_path), download_url
